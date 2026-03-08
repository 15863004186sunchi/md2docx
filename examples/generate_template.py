from __future__ import annotations

from pathlib import Path

from docx import Document


def main() -> None:
    out = Path(__file__).parent / "template.docx"
    doc = Document()
    doc.add_heading("项目背景", level=1)
    doc.add_paragraph("（占位内容：将被替换）")

    doc.add_heading("业务目标", level=2)
    doc.add_paragraph("（占位内容：将被替换）")

    doc.add_heading("核心流程", level=2)
    doc.add_paragraph("（占位内容：将被替换）")

    doc.add_heading("风险与应对", level=2)
    doc.add_paragraph("（占位内容：将被替换）")

    doc.save(out)
    print(f"template generated: {out}")


if __name__ == "__main__":
    main()

