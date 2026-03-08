from __future__ import annotations

from pathlib import Path

from docx.document import Document as DocxDocument
from docx.oxml import OxmlElement
from docx.shared import Inches
from docx.text.paragraph import Paragraph

from .mermaid_renderer import MermaidRenderer
from .models import (
    CodeBlock,
    ConversionReport,
    InsertConfig,
    ListBlock,
    MatchResult,
    MermaidBlock,
    ParagraphBlock,
    SectionNode,
    TemplateHeading,
)


def apply_mappings(
    doc: DocxDocument,
    sections_by_path: dict[str, SectionNode],
    matches: list[MatchResult],
    headings: list[TemplateHeading],
    insert_cfg: InsertConfig,
    mermaid_renderer: MermaidRenderer,
    report: ConversionReport,
    image_ratio: float,
) -> None:
    heading_map = {heading.paragraph_index: heading for heading in headings}
    boundary_by_idx = _build_boundary_map(headings, len(doc.paragraphs))

    valid_matches = [x for x in matches if x.template_paragraph_index is not None]
    for match in sorted(valid_matches, key=lambda item: item.template_paragraph_index or -1, reverse=True):
        section = sections_by_path.get(match.md_path)
        if section is None:
            report.warnings.append(f"Section not found for mapping path: {match.md_path}")
            continue

        heading_idx = match.template_paragraph_index
        if heading_idx is None or heading_idx >= len(doc.paragraphs):
            report.warnings.append(f"Invalid heading index for path: {match.md_path}")
            continue

        if insert_cfg.mode == "replace":
            boundary = boundary_by_idx.get(heading_idx, len(doc.paragraphs))
            _delete_range(doc, heading_idx + 1, min(boundary, len(doc.paragraphs)))

        anchor = doc.paragraphs[heading_idx]
        for block in section.blocks:
            anchor = _insert_block(
                anchor=anchor,
                block=block,
                insert_cfg=insert_cfg,
                doc=doc,
                mermaid_renderer=mermaid_renderer,
                report=report,
                image_ratio=image_ratio,
            )

        if heading_idx in heading_map and not section.blocks:
            report.warnings.append(f"Section '{section.title_raw}' is empty in markdown")


def _build_boundary_map(headings: list[TemplateHeading], total: int) -> dict[int, int]:
    by_index = sorted(headings, key=lambda h: h.paragraph_index)
    result: dict[int, int] = {}
    for i, heading in enumerate(by_index):
        next_idx = by_index[i + 1].paragraph_index if i + 1 < len(by_index) else total
        result[heading.paragraph_index] = next_idx
    return result


def _delete_range(doc: DocxDocument, start: int, end: int) -> None:
    if start >= end:
        return
    for idx in range(end - 1, start - 1, -1):
        paragraph = doc.paragraphs[idx]
        element = paragraph._element
        element.getparent().remove(element)


def _insert_block(
    anchor: Paragraph,
    block: ParagraphBlock | ListBlock | CodeBlock | MermaidBlock,
    insert_cfg: InsertConfig,
    doc: DocxDocument,
    mermaid_renderer: MermaidRenderer,
    report: ConversionReport,
    image_ratio: float,
) -> Paragraph:
    if isinstance(block, ParagraphBlock):
        return _insert_plain_paragraph(anchor, block.text)
    if isinstance(block, ListBlock):
        current = anchor
        for item in block.items:
            marker = "1." if block.ordered else insert_cfg.list_fallback_prefix
            current = _insert_list_paragraph(current, f"{marker} {item}".strip(), ordered=block.ordered)
        return current
    if isinstance(block, CodeBlock):
        para = _insert_plain_paragraph(anchor, block.code)
        for run in para.runs:
            run.font.name = insert_cfg.code_font_name
        return para
    return _insert_mermaid_block(anchor, block, doc, mermaid_renderer, report, image_ratio)


def _insert_plain_paragraph(anchor: Paragraph, text: str) -> Paragraph:
    paragraph = _insert_paragraph_after(anchor, text=text)
    _apply_style_if_exists(paragraph, "Normal")
    return paragraph


def _insert_list_paragraph(anchor: Paragraph, text: str, ordered: bool) -> Paragraph:
    paragraph = _insert_paragraph_after(anchor, text=text)
    desired = "List Number" if ordered else "List Bullet"
    if not _apply_style_if_exists(paragraph, desired):
        _apply_style_if_exists(paragraph, "Normal")
    return paragraph


def _insert_mermaid_block(
    anchor: Paragraph,
    block: MermaidBlock,
    doc: DocxDocument,
    mermaid_renderer: MermaidRenderer,
    report: ConversionReport,
    image_ratio: float,
) -> Paragraph:
    rendered = mermaid_renderer.render(block.code)
    if not rendered.ok or rendered.image_path is None:
        message = rendered.error or "Unknown mermaid error"
        report.mermaid_failures.append(message)
        return _insert_plain_paragraph(anchor, f"[Mermaid 渲染失败] {message}\n{block.code}")

    para = _insert_paragraph_after(anchor)
    run = para.add_run()
    run.add_picture(str(rendered.image_path), width=Inches(_available_width_inches(doc) * image_ratio))
    _apply_style_if_exists(para, "Normal")
    return para


def _available_width_inches(doc: DocxDocument) -> float:
    section = doc.sections[0]
    available = section.page_width - section.left_margin - section.right_margin
    if hasattr(available, "inches"):
        return float(available.inches)
    return float(available) / 914400.0


def _insert_paragraph_after(anchor: Paragraph, text: str = "") -> Paragraph:
    new_p = OxmlElement("w:p")
    anchor._p.addnext(new_p)
    new_para = Paragraph(new_p, anchor._parent)
    if text:
        new_para.add_run(text)
    return new_para


def _apply_style_if_exists(paragraph: Paragraph, style_name: str) -> bool:
    style_names = {style.name for style in paragraph.part.document.styles}
    if style_name in style_names:
        paragraph.style = style_name
        return True
    return False
